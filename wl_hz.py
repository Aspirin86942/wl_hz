import xlrd
from docx import Document
import re
import os
import shutil
from comtypes.client import CreateObject


# 从excel取数到字典中处理
class Excel2Dict(object):
    def __init__(self, input_path, sheet_name):
        self.input_path = input_path
        self.sheet_name = sheet_name

    def read_excel(self):
        # 打开excel表，填写路径
        book = xlrd.open_workbook(self.input_path)
        # 找到sheet页
        table = book.sheet_by_name(self.sheet_name)
        # 获取总行数总列数
        row_num = table.nrows
        col_num = table.ncols
        ls = []
        key = table.row_values(0)
        # 这是第一行数据，作为字典的key值
        if row_num <= 1:
            # print("没数据")
            return ls
        else:
            j = 1
            for i in range(row_num - 1):
                d = {}
                values = table.row_values(j)
                for x in range(col_num):
                    # 把key值对应的value赋值给key，每行循环
                    d[key[x]] = values[x]
                j += 1
                # 把字典加到列表中
                ls.append(d)
            return ls

    def distinct_key(self, dict):
        main_key = []
        catalogue_dict = {}
        for i in range(len(dict)):
            main_key.append(dict[i]["询证函编号"])
        for item in main_key:
            catalogue_dict.update({item: ""})
        keys = []
        for key in catalogue_dict:
            keys.append(key)
        return keys

    def distinct_list(self, key, dict):
        table = []
        for j in range(len(dict)):
            if dict[j]["询证函编号"] == key:
                table.append(dict[j])
        return table


class Dict2Word(object):
    def __init__(self, key, inf_table):
        self.key = key
        self.inf_table = inf_table

    def make_doc(self):
        global replace_dict  # 设定基本信息字典
        replace_dict = {}
        document = Document('往来函证模板.docx')

        # 设定表
        balance_table = document.tables[1]
        # 比较一下所填内容和word中银行存款的表格大小
        if len(balance_table.rows) - 1 <= len(self.inf_table):
            # 增加行
            for j in range(len(self.inf_table) - len(balance_table.rows) + 1):
                balance_table.add_row()
        for i in range(1, len(balance_table.rows)):
            date = \
                xlrd.xldate.xldate_as_datetime(self.inf_table[i - 1]["结算日期"], 0).strftime('%Y-%m-%d')
            balance_table.cell(i, 0).text = str(date)
            balance_table.cell(i, 1).text = reg_balance(self.inf_table[i - 1]["贵公司欠本公司"])
            balance_table.cell(i, 2).text = reg_balance(self.inf_table[i - 1]["本公司欠贵公司"])
            balance_table.cell(i, 3).text = self.inf_table[i - 1]["备注"]
        # 函证基本信息填充部分
        replace_dict.update({
            "编码": self.inf_table[0]["询证函编号"],
            # "数据2": self.inf_table[0]["函证日期"],
            "对方": self.inf_table[0]["客户或供应商名称"],
            "我方单位": self.inf_table[0]["发函单位"]
        })
        # 开始替换
        hz_replace(document=document, replace_dict=replace_dict)
        # 保存函证
        document.save(r'.\word版函证\%s.docx' % self.key)
        # 进度提示
        print("%s生成完毕" % self.key)


# 静态方法集合
# 替换函证信息的方法
def hz_replace(document, replace_dict):

    for table in document.tables:
        for row in range(len(table.rows)):
            for col in range(len(table.columns)):
                for key, value in replace_dict.items():
                    if key in table.cell(row, col).text:
                        # print(key + "->" + value)
                        table.cell(row, col).text = table.cell(row, col).text.replace(key, value)

    for para in document.paragraphs:
        for i in range(len(para.runs)):
            for key, value in replace_dict.items():
                if key in para.runs[i].text:
                    para.runs[i].text = para.runs[i].text.replace(key, value)
                # print(para.runs[i].text)


# 创建word文件夹
def setup_doc_file():
    is_exists = os.path.exists(r".\word版函证")
    if not is_exists:
        os.makedirs(r".\word版函证")
    else:
        shutil.rmtree(r".\word版函证")
        os.makedirs(r".\word版函证")


# 创建pdf文件夹
def setup_pdf_file():
    is_exists = os.path.exists(r".\pdf版函证")
    if not is_exists:
        os.makedirs(r".\pdf版函证")
    else:
        shutil.rmtree(r".\pdf版函证")
        os.makedirs(r".\pdf版函证")


# 对于余额浮点数的规范
def reg_balance(balance):
    if balance == "":
        return balance
    else:
        balance_1 = '%.2f' % (float(balance))
        balance_format_1 = re.sub(r"(\d)(?=(\d\d\d)+(?!\d))", r"\1,", balance_1)
        return str(balance_format_1)


class Word2Pdf(object):
    def __init__(self):
        # word文档转化为pdf文档时使用的格式为17
        self.wdFormatPDF = 17
        self.wdToPDF = CreateObject("Word.Application")

    def wd_to_pdf(self, input_path, output_path):
        # 获取指定目录下面的所有文件
        files = os.listdir(input_path)
        # files = list_nohidden(folder)
        # print(files)
        # 获取word类型的文件放到一个列表里面
        wdfiles = [f for f in files if f.endswith((".doc", ".docx"))]
        # 去除word生成的隐藏文件
        wdfiles2 = [f for f in wdfiles if not f.startswith('~')]
        for wdfile in wdfiles2:
            # 将word文件放到指定的路径下面
            wdPath = os.path.join(input_path, wdfile)
            # print(wdfile)
            # 设置将要存放pdf文件的路径
            pdfPath = output_path + wdfile.split(".")[0] + '.pdf'
            # 判断是否已经存在对应的pdf文件，如果不存在就加入到存放pdf的路径内
            if pdfPath[-3:] != 'pdf':
                pdfPath = pdfPath + ".pdf"
            # 将word文档转化为pdf文件，先打开word所在路径文件，然后在处理后保存pdf文件，最后关闭
            pdfCreate = self.wdToPDF.Documents.Open(wdPath)
            pdfCreate.SaveAs(pdfPath, self.wdFormatPDF)
            pdfCreate.Close()
            print("正在生成%s" % pdfPath)


def main():
    # 设定工作目录
    work_path = os.getcwd()
    os.chdir(work_path)

    # 创建文件夹
    setup_doc_file()
    setup_pdf_file()

    # 实例化工作簿转字典类
    wl_excel_dict = Excel2Dict(input_path="wl_hz.xlsx", sheet_name="Sheet1")
    wl_dict = wl_excel_dict.read_excel()
    # print(wl_dict)
    # 对于一般有重复值的函证清单，根据询证函编号做出唯一值发函清单
    wl_dist_key = wl_excel_dict.distinct_key(dict=wl_dict)
    # print(wl_dist_key)
    # 对于每一份函证单独生成一个table
    for key in wl_dist_key:
        single_sl_table = wl_excel_dict.distinct_list(key=key, dict=wl_dict)
        # 实例化字典转doc类
        wl_dict_doc = Dict2Word(key=key, inf_table=single_sl_table)
        # 报错信息
        try:
            wl_dict_doc.make_doc()
        except ValueError:
            print("ValueError,%s不生成" % key)
        except IndexError:
            print("IndexError,%s不生成" % key)

    # 生成函证后自动生成pdf
    print("正在生成pdf")
    converter = Word2Pdf()
    input_path = work_path + r"\word版函证"
    output_path = work_path + r"\pdf版函证\\"
    converter.wd_to_pdf(input_path=input_path, output_path=output_path)
    print("pdf生成完毕")
    # 保证窗口在程序结束后不消失
    input('Press Enter to exit...')


if __name__ == '__main__':
    main()
